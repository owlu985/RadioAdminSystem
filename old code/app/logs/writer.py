from docx import Document
import os

def csv_to_docx(csv_path):
    doc = Document()
    doc.add_heading("WLMC Radio Log", level=1)

    with open(csv_path) as f:
        for line in f:
            doc.add_paragraph(line.strip())

    out_path = csv_path.replace(".csv", ".docx")
    doc.save(out_path)

    return out_path
