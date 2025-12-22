from flask import Blueprint, render_template, request, send_file, current_app
from sqlalchemy import asc, desc
from io import BytesIO
import csv
from docx import Document

from .models import db, LogEntry, Show
from .routes import admin_required
from .logger import init_logger

logs_admin_bp = Blueprint('logs_admin', __name__, url_prefix='/logs')
logger = init_logger()

@logs_admin_bp.route('/')
@admin_required
def logs_index():
	"""
	Admin log browser with sorting.
	"""
	sort = request.args.get('sort', 'date')
	order = request.args.get('order', 'desc')

	query = LogEntry.query.join(Show, isouter=True)

	if sort == 'dj':
		column = Show.host_last_name
	else:
		column = LogEntry.timestamp

	query = query.order_by(desc(column) if order == 'desc' else asc(column))
	logs = query.limit(500).all()  # safety cap

	return render_template(
		'logs_index.html',
		logs=logs,
		sort=sort,
		order=order
	)


@logs_admin_bp.route('/view/<int:log_id>')
@admin_required
def view_log(log_id):
	"""
	View a single log entry.
	"""
	log = LogEntry.query.get_or_404(log_id)
	return render_template('log_view.html', log=log)
@logs_admin_bp.route('/export/csv')
@admin_required
def export_csv():
	logs = LogEntry.query.order_by(LogEntry.timestamp).all()

	output = BytesIO()
	writer = csv.writer(output)

	writer.writerow([
		"Timestamp", "Type", "Title", "Artist",
		"Description", "Show", "DJ"
	])

	for log in logs:
		show = log.show
		writer.writerow([
			log.timestamp,
			log.entry_type,
			log.title or "",
			log.artist or "",
			log.description or "",
			show.show_name if show else "",
			f"{show.host_first_name} {show.host_last_name}" if show else ""
		])

	output.seek(0)
	return send_file(
		output,
		mimetype='text/csv',
		as_attachment=True,
		download_name='wlmc_logs.csv'
	)


@logs_admin_bp.route('/export/docx')
@admin_required
def export_docx():
	logs = LogEntry.query.order_by(LogEntry.timestamp).all()

	doc = Document()
	doc.add_heading('WLMC Station Logs', level=1)

	for log in logs:
		show = log.show
		p = doc.add_paragraph()
		p.add_run(f"{log.timestamp} — {log.entry_type.upper()}\n").bold = True
		if log.title:
			p.add_run(f"Title: {log.title}\n")
		if log.artist:
			p.add_run(f"Artist: {log.artist}\n")
		if log.description:
			p.add_run(f"Notes: {log.description}\n")
		if show:
			p.add_run(
				f"Show: {show.show_name} "
				f"({show.host_first_name} {show.host_last_name})\n"
			)

	output = BytesIO()
	doc.save(output)
	output.seek(0)

	return send_file(
		output,
		as_attachment=True,
		download_name='wlmc_logs.docx'
	)
